import os
from docx import Document

#=== CONFIGURATION ===
# Set this to the folder you want to analyze
base_path = '/yor directory'  # or '.' if already in folder
output_docx = 'folder_structure.docx'

#=== STEP 1: Create Word Document ===
doc = Document()
doc.add_heading('Folder and File Structure', level=1)

#=== STEP 2: Walk Through Folder ===
def add_folder_contents(doc, path, indent=0):
    prefix = '    ' * indent + '• '
    items = sorted(os.listdir(path))
    for item in items:
        item_path = os.path.join(path, item)
        doc.add_paragraph(f"{prefix}{item}")
        if os.path.isdir(item_path):
            add_folder_contents(doc, item_path, indent + 1)

add_folder_contents(doc, base_path)

#=== STEP 3: Save Word File ===
doc.save(output_docx)
print(f"Saved as {output_docx}")
